#!/usr/bin/env python3
"""report_docx.py — assemble the stage-7 figure report as a Word (.docx) document.

Builds the hand-off report: a title, manuscript metadata, an asset index, then one
block per figure/table with the embedded figure image (or inline three-line table),
caption, annotations, in-text citation location, and a reproducibility note.

Output language is selectable via `lang`:
  - "en"        — English only
  - "zh"        — Chinese only / 全中文
  - "bilingual" — both available English and Chinese text variants (default)

    from report_docx import build_report
    build_report("Figure_Report.docx", title=..., meta=..., items=..., lang="en")

Each item is a dict. Text fields accept either a plain string or a
``{"en": "...", "zh": "..."}`` dict; when a string is given it is used for every
language (handy when the text is language-neutral). Legacy ``caption_en`` /
``caption_zh`` keys are still accepted.

  figure: {"kind":"figure", "number":"1", "title": {"en":..,"zh":..},
           "image":"figures/Fig1.png", "width_in":6.0,
           "claim": {"en":..,"zh":..}, "caption": {"en":..,"zh":..},
           "annotations": {"en":[...], "zh":[...]},   # or a shared list
           "method_rationale": {...}, "result": {...}, "interpretation": {...},
           "how_to_read": {"en":[...], "zh":[...]}, "common_misreading": {...},
           "cannot_conclude": {...},
           "status": "draft|needs-author-decision|approved|final",
           "citation": {"en":..,"zh":..}, "repro":"scripts/make_fig1.py"}
  table:  {"kind":"table", "number":"1", "title": {...}, "csv":"figures/Table1.csv",
           "table_title": {...}, "table_note": {...}, ...same text fields... }

Requires python-docx (and pandas for table CSVs).
"""
from __future__ import annotations

import os
from pathlib import Path

from docx_tables import add_three_line_table

# section labels per language
_LABELS = {
    "bilingual": {"asset": "Asset index / 图表清单", "cols": ["#", "File 文件", "Claim 结论"],
                  "claim": "Claim / 论证: ", "cap_en": "Caption (EN): ", "cap_zh": "图注 (中文): ",
                  "ann": "Annotations / 标注", "cite": "Citation location / 引用位置: ",
                  "repro": "Reproducibility / 可复现: ", "fig": "Figure ", "tbl": "Table "},
    "en": {"asset": "Asset index", "cols": ["#", "File", "Claim"],
           "claim": "Claim: ", "cap": "Caption: ", "ann": "Annotations",
           "cite": "Citation location: ", "repro": "Reproducibility: ",
           "fig": "Figure ", "tbl": "Table "},
    "zh": {"asset": "图表清单", "cols": ["#", "文件", "结论"],
           "claim": "论证:", "cap": "图注:", "ann": "标注",
           "cite": "引用位置:", "repro": "可复现:", "fig": "图 ", "tbl": "表 "},
}

_EXTRA_LABELS = {
    "bilingual": {
        "method": "Method rationale / 方法依据: ", "result": "Scientific result / 科研结果: ",
        "interpretation": "Interpretation / 结果解读: ", "reading": "How to read / 阅读方法",
        "misreading": "Common misreading / 常见误读: ",
        "cannot": "Cannot conclude / 不能据此得出: ", "status": "Review status / 审核状态: ",
    },
    "en": {
        "method": "Method rationale: ", "result": "Scientific result: ",
        "interpretation": "Interpretation: ", "reading": "How to read",
        "misreading": "Common misreading: ", "cannot": "Cannot conclude: ",
        "status": "Review status: ",
    },
    "zh": {
        "method": "方法依据:", "result": "科研结果:", "interpretation": "结果解读:",
        "reading": "阅读方法", "misreading": "常见误读:", "cannot": "不能据此得出:",
        "status": "审核状态:",
    },
}

_FIELD_LABELS = {
    "claim": {"en": "Claim: ", "zh": "论证：", "both": "Claim / 论证: "},
    "caption": {"en": "Caption: ", "zh": "图注：", "both": "Caption / 图注: "},
    "annotations": {"en": "Annotations", "zh": "标注", "both": "Annotations / 标注"},
    "citation": {"en": "Citation location: ", "zh": "引用位置：", "both": "Citation location / 引用位置: "},
    "repro": {"en": "Reproducibility: ", "zh": "可复现：", "both": "Reproducibility / 可复现: "},
    "method": {"en": "Method rationale: ", "zh": "方法依据：", "both": "Method rationale / 方法依据: "},
    "result": {"en": "Scientific result: ", "zh": "科研结果：", "both": "Scientific result / 科研结果: "},
    "interpretation": {"en": "Interpretation: ", "zh": "结果解读：", "both": "Interpretation / 结果解读: "},
    "reading": {"en": "How to read", "zh": "阅读方法", "both": "How to read / 阅读方法"},
    "misreading": {"en": "Common misreading: ", "zh": "常见误读：", "both": "Common misreading / 常见误读: "},
    "cannot": {"en": "Cannot conclude: ", "zh": "不能据此得出：", "both": "Cannot conclude / 不能据此得出: "},
    "status": {"en": "Review status: ", "zh": "审核状态：", "both": "Review status / 审核状态: "},
}

_VALID_KINDS = {"figure", "table"}
_VALID_STATUSES = {"draft", "needs-author-decision", "approved", "final"}


def _pick(val, lang):
    """Resolve a string-or-{en,zh} field to one string for `lang`."""
    if isinstance(val, dict):
        if lang == "zh":
            return val.get("zh") or val.get("en") or ""
        return val.get("en") or val.get("zh") or ""
    return val or ""


def _present(value):
    """Return whether a localized value contains displayable content."""
    if value is None:
        return False
    if isinstance(value, (str, bytes, list, tuple, dict)):
        return bool(value)
    return True


def _localized_entries(value, lang):
    """Resolve a shared or ``{en, zh}`` value without dropping bilingual text.

    The first tuple element is ``en``/``zh`` when the source explicitly carries a
    language, and ``None`` for a shared value. In a monolingual report, the old
    fallback behaviour (requested language, then the other language) is retained.
    """
    if isinstance(value, dict):
        if lang == "bilingual":
            return [(code, value.get(code)) for code in ("en", "zh")
                    if _present(value.get(code))]
        selected = _pick(value, lang)
        return [(lang, selected)] if _present(selected) else []
    return [(None, value)] if _present(value) else []


def _localized_join(value, lang, separator="\n"):
    """Return all applicable localized strings joined for one Word run."""
    return separator.join(str(text) for _, text in _localized_entries(value, lang))


def _field_label(field, lang, source_lang=None):
    labels = _FIELD_LABELS[field]
    if lang == "bilingual":
        return labels.get(source_lang, labels["both"])
    return labels["zh" if lang == "zh" else "en"]


def _render_text_field(doc, field, value, lang, size=10):
    """Render every applicable language variant of a scalar text field."""
    paragraphs = []
    for source_lang, text in _localized_entries(value, lang):
        paragraphs.append(
            _labelled(doc, _field_label(field, lang, source_lang), str(text), size=size)
        )
    return paragraphs


def _render_bullet_field(doc, field, value, lang, size=10, annotations=False):
    """Render localized lists, preserving both language lists in bilingual mode."""
    for source_lang, values in _localized_entries(value, lang):
        heading = _labelled(doc, _field_label(field, lang, source_lang), "", size=size)
        heading.paragraph_format.keep_with_next = True
        if isinstance(values, str):
            values = [values]
        for entry in values or []:
            if annotations and isinstance(entry, (tuple, list)):
                if len(entry) != 2:
                    raise ValueError(
                        f"annotation pairs must contain exactly two values, got {entry!r}"
                    )
                from docx.shared import Pt
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(f"{entry[0]}: ")
                r.bold = True
                r.font.size = Pt(size)
                p.add_run(str(entry[1])).font.size = Pt(size)
            else:
                _bullets(doc, [entry], size=size)


def _meta_rows(label, value, lang):
    """Pair localized metadata labels and values for display."""
    if lang != "bilingual" or not isinstance(label, dict) and not isinstance(value, dict):
        chosen = "zh" if lang == "zh" else "en"
        return [(_pick(label, chosen), _pick(value, chosen))]

    rows = []
    for code in ("en", "zh"):
        label_explicit = isinstance(label, dict) and _present(label.get(code))
        value_explicit = isinstance(value, dict) and _present(value.get(code))
        if not (label_explicit or value_explicit):
            continue
        row_label = _pick(label, code) if isinstance(label, dict) else label
        row_value = _pick(value, code) if isinstance(value, dict) else value
        rows.append((row_label, row_value))
    return rows


def _asset_path(base, item):
    key = "image" if item["kind"] == "figure" else "csv"
    raw = item[key]
    candidate = Path(raw)
    return candidate if candidate.is_absolute() else base / candidate


def _validate_items(items, base):
    """Validate the small public item schema and every referenced input asset."""
    if items is None:
        return []
    if isinstance(items, (str, bytes, dict)):
        raise TypeError("items must be an iterable of item dictionaries")
    try:
        normalized = list(items)
    except TypeError as exc:
        raise TypeError("items must be an iterable of item dictionaries") from exc

    seen = set()
    for index, item in enumerate(normalized):
        where = f"items[{index}]"
        if not isinstance(item, dict):
            raise TypeError(f"{where} must be a dictionary, got {type(item).__name__}")

        kind = item.get("kind")
        if not isinstance(kind, str) or kind not in _VALID_KINDS:
            raise ValueError(f"{where}.kind must be 'figure' or 'table', got {kind!r}")

        number = item.get("number")
        if isinstance(number, bool) or not isinstance(number, (str, int)) or not str(number).strip():
            raise ValueError(f"{where}.number must be a non-empty string or integer")
        identity = (kind, str(number))
        if identity in seen:
            raise ValueError(f"duplicate {kind} number {number!r}")
        seen.add(identity)

        status = item.get("status")
        if status is not None and (not isinstance(status, str) or status not in _VALID_STATUSES):
            allowed = ", ".join(sorted(_VALID_STATUSES))
            raise ValueError(f"{where}.status must be one of {allowed}, got {status!r}")

        asset_key = "image" if kind == "figure" else "csv"
        asset = item.get(asset_key)
        if not isinstance(asset, (str, os.PathLike)) or not str(asset).strip():
            raise ValueError(f"{where}.{asset_key} is required for every {kind}")
        asset_path = _asset_path(base, item)
        if not asset_path.is_file():
            raise FileNotFoundError(
                f"{kind.capitalize()} {number} {asset_key} not found: {asset_path}"
            )

        if kind == "figure" and "width_in" in item:
            width = item["width_in"]
            if isinstance(width, bool) or not isinstance(width, (int, float)) or width <= 0:
                raise ValueError(f"{where}.width_in must be a positive number")
    return normalized


def _caption_dict(it):
    """Normalize caption to {en, zh}, supporting legacy caption_en/caption_zh."""
    cap = it.get("caption")
    if isinstance(cap, dict):
        return cap
    if isinstance(cap, str):
        return {"en": cap, "zh": cap}
    return {"en": it.get("caption_en", ""), "zh": it.get("caption_zh", "")}


def _caption_value(it):
    """Return a caption while preserving whether a modern string is shared."""
    caption = it.get("caption")
    if isinstance(caption, (str, dict)):
        return caption
    return {"en": it.get("caption_en", ""), "zh": it.get("caption_zh", "")}


def _item_reference(kind, number, lang, source_lang=None):
    prefixes = {"figure": {"en": "Figure ", "zh": "图 "},
                "table": {"en": "Table ", "zh": "表 "}}
    if lang == "bilingual":
        if source_lang in ("en", "zh"):
            return f"{prefixes[kind][source_lang]}{number}"
        return f"{prefixes[kind]['en']}{number} / {prefixes[kind]['zh']}{number}"
    code = "zh" if lang == "zh" else "en"
    return f"{prefixes[kind][code]}{number}"


def _item_heading(item, lang):
    """Build an item heading containing both localized titles when requested."""
    title_entries = _localized_entries(item.get("title"), lang)
    if not title_entries:
        return _item_reference(item["kind"], item["number"], lang)
    parts = []
    for source_lang, title in title_entries:
        reference = _item_reference(item["kind"], item["number"], lang, source_lang)
        parts.append(f"{reference} — {title}")
    return "\n".join(parts)


def _heading(doc, text, size, bold=True, space_before=10, space_after=4, color=None,
             leading_spacer=False):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if leading_spacer:
        spacer = p.add_run("MMMMMMMM ")
        spacer.font.size = Pt(size)
        spacer.font.color.rgb = RGBColor(255, 255, 255)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _labelled(doc, label, text, size=10, label_color=(0x1F, 0x4E, 0x79)):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*label_color)
    if text:
        t = p.add_run(text)
        t.font.size = Pt(size)
    return p


def _bullets(doc, values, size=10):
    from docx.shared import Pt
    if isinstance(values, str):
        values = [values]
    for value in values or []:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(str(value)).font.size = Pt(size)


def build_report(path, title, subtitle=None, meta=None, items=None,
                 base_dir=None, lang="bilingual", body_font="Calibri", body_pt=10):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if lang not in _LABELS:
        raise ValueError(f"lang must be one of {list(_LABELS)}, got {lang!r}")
    L = _LABELS[lang]
    base = Path(base_dir) if base_dir else Path(".")
    items = _validate_items(items, base)
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = body_font
    s.font.size = Pt(body_pt)

    title_entries = _localized_entries(title, lang)
    if not title_entries:
        raise ValueError("title must contain displayable text")
    for index, (_, text) in enumerate(title_entries):
        _heading(
            doc, str(text), 18, color=(0x1F, 0x4E, 0x79), space_before=0,
            space_after=2 if index == len(title_entries) - 1 else 0,
        )
    if subtitle:
        subtitle_entries = _localized_entries(subtitle, lang)
        for index, (_, text) in enumerate(subtitle_entries):
            _heading(
                doc, str(text), 12, bold=False, space_before=0,
                space_after=8 if index == len(subtitle_entries) - 1 else 0,
            )

    if meta:
        for index, entry in enumerate(meta):
            if not isinstance(entry, (tuple, list)) or len(entry) != 2:
                raise ValueError(f"meta[{index}] must be a (label, value) pair")
            label, value = entry
            for row_label, row_value in _meta_rows(label, value, lang):
                _labelled(doc, f"{row_label}: ", str(row_value), size=body_pt)

    if items:
        _heading(doc, L["asset"], 12, space_before=12)
        idx = doc.add_table(rows=1, cols=3)
        idx.style = "Light Grid Accent 1"
        for j, h in enumerate(L["cols"]):
            r = idx.rows[0].cells[j].paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(body_pt - 1)
        for it in items:
            cells = idx.add_row().cells
            fname = Path(it["image"] if it["kind"] == "figure" else it["csv"]).name
            values = [
                _item_reference(it["kind"], it["number"], lang),
                fname,
                _localized_join(it.get("claim"), lang),
            ]
            for j, value in enumerate(values):
                rr = cells[j].paragraphs[0].add_run(str(value))
                rr.font.size = Pt(body_pt - 1)

    for it in items:
        if it.get("page_break_before"):
            from docx.shared import RGBColor
            break_paragraph = doc.add_paragraph()
            break_paragraph.paragraph_format.page_break_before = True
            break_paragraph.paragraph_format.space_before = Pt(0)
            break_paragraph.paragraph_format.space_after = Pt(0)
            marker = break_paragraph.add_run(".")
            marker.font.size = Pt(1)
            marker.font.color.rgb = RGBColor(255, 255, 255)
        head = f"{it.get('heading_guard', '')}{_item_heading(it, lang)}"
        if it["kind"] == "table":
            head_paragraph = _labelled(doc, head, "", size=13)
            head_paragraph.paragraph_format.space_before = Pt(14)
            head_paragraph.paragraph_format.space_after = Pt(4)
        else:
            head_paragraph = _heading(
                doc, head, 13, color=(0x1F, 0x4E, 0x79), space_before=14,
            )

        if it["kind"] == "figure" and not it.get("image_after_caption"):
            img = _asset_path(base, it)
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.add_run().add_picture(str(img), width=Inches(it.get("width_in", 6.0)))
            except Exception as exc:
                raise ValueError(f"Could not embed image for Figure {it['number']}: {img}") from exc
        elif it["kind"] == "table":
            import pandas as pd
            csv_path = _asset_path(base, it)
            try:
                df = pd.read_csv(csv_path)
            except Exception as exc:
                raise ValueError(f"Could not read CSV for Table {it['number']}: {csv_path}") from exc
            if df.shape[1] == 0:
                raise ValueError(f"CSV for Table {it['number']} has no columns: {csv_path}")
            add_three_line_table(doc, df,
                                 title=_localized_join(it.get("table_title"), lang) or None,
                                 note=_localized_join(it.get("table_note"), lang) or None,
                                 size_pt=max(7, body_pt - 1))

        if it.get("claim"):
            _render_text_field(doc, "claim", it["claim"], lang, size=body_pt)
        if it.get("method_rationale"):
            _render_text_field(doc, "method", it["method_rationale"], lang, size=body_pt)
        if it.get("result"):
            _render_text_field(doc, "result", it["result"], lang, size=body_pt)
        if it.get("interpretation"):
            _render_text_field(doc, "interpretation", it["interpretation"], lang, size=body_pt)

        cap = _caption_value(it)
        if _present(cap):
            _render_text_field(doc, "caption", cap, lang, size=body_pt)

        if it["kind"] == "figure" and it.get("image_after_caption"):
            img = _asset_path(base, it)
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.add_run().add_picture(str(img), width=Inches(it.get("width_in", 6.0)))
            except Exception as exc:
                raise ValueError(f"Could not embed image for Figure {it['number']}: {img}") from exc

        anns = it.get("annotations")
        if anns:
            _render_bullet_field(
                doc, "annotations", anns, lang, size=body_pt, annotations=True,
            )

        reading = it.get("how_to_read")
        if reading:
            _render_bullet_field(doc, "reading", reading, lang, size=body_pt)
        if it.get("common_misreading"):
            _render_text_field(doc, "misreading", it["common_misreading"], lang, size=body_pt)
        if it.get("cannot_conclude"):
            _render_text_field(doc, "cannot", it["cannot_conclude"], lang, size=body_pt)
        if it.get("status"):
            _render_text_field(doc, "status", it["status"], lang, size=body_pt)

        if it.get("citation"):
            _render_text_field(doc, "citation", it["citation"], lang, size=body_pt)
        if it.get("repro"):
            _render_text_field(doc, "repro", it["repro"], lang, size=body_pt)

    doc.save(str(output))
    return path
