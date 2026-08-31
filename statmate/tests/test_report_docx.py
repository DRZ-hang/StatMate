"""Regression tests for the StatMate Word report builder.

The suite uses only the standard-library ``unittest`` runner so it can run in the
bundled workspace Python without adding a pytest dependency.
"""
from __future__ import annotations

import base64
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS))

from report_docx import build_report  # noqa: E402


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
    "AQUBAScY42YAAAAASUVORK5CYII="
)


def _document_text(path: Path) -> str:
    """Collect visible text from document paragraphs and top-level table cells."""
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


class ReportDocxTests(unittest.TestCase):
    def setUp(self):
        self._temporary = tempfile.TemporaryDirectory()
        self.root = Path(self._temporary.name)
        self.image = self.root / "figure.png"
        self.image.write_bytes(_ONE_PIXEL_PNG)
        self.csv = self.root / "table.csv"
        self.csv.write_text("Measure,Value\nAlpha,1\n", encoding="utf-8")

    def tearDown(self):
        self._temporary.cleanup()

    def _figure(self, **updates):
        item = {
            "kind": "figure",
            "number": "1",
            "image": self.image.name,
            "caption": "Shared caption",
        }
        item.update(updates)
        return item

    def _table(self, **updates):
        item = {
            "kind": "table",
            "number": "1",
            "csv": self.csv.name,
        }
        item.update(updates)
        return item

    def test_bilingual_renders_both_languages_for_every_localized_field(self):
        output = self.root / "new" / "nested" / "report.docx"
        figure = self._figure(
            title={"en": "FIGURE_TITLE_EN", "zh": "FIGURE_TITLE_ZH"},
            claim={"en": "CLAIM_EN", "zh": "CLAIM_ZH"},
            method_rationale={"en": "METHOD_EN", "zh": "METHOD_ZH"},
            result={"en": "RESULT_EN", "zh": "RESULT_ZH"},
            interpretation={"en": "INTERPRET_EN", "zh": "INTERPRET_ZH"},
            caption={"en": "CAPTION_EN", "zh": "CAPTION_ZH"},
            annotations={
                "en": [("ANNOTATION_KEY_EN", "ANNOTATION_VALUE_EN")],
                "zh": [("ANNOTATION_KEY_ZH", "ANNOTATION_VALUE_ZH")],
            },
            how_to_read={"en": ["READ_EN"], "zh": ["READ_ZH"]},
            common_misreading={"en": "MISREAD_EN", "zh": "MISREAD_ZH"},
            cannot_conclude={"en": "LIMIT_EN", "zh": "LIMIT_ZH"},
            citation={"en": "CITATION_EN", "zh": "CITATION_ZH"},
            repro={"en": "REPRO_EN", "zh": "REPRO_ZH"},
            status="final",
        )
        table = self._table(
            number="2",
            title={"en": "TABLE_HEADING_EN", "zh": "TABLE_HEADING_ZH"},
            table_title={"en": "TABLE_TITLE_EN", "zh": "TABLE_TITLE_ZH"},
            table_note={"en": "TABLE_NOTE_EN", "zh": "TABLE_NOTE_ZH"},
        )

        returned = build_report(
            output,
            title={"en": "REPORT_TITLE_EN", "zh": "REPORT_TITLE_ZH"},
            subtitle={"en": "SUBTITLE_EN", "zh": "SUBTITLE_ZH"},
            meta=[
                (
                    {"en": "META_LABEL_EN", "zh": "META_LABEL_ZH"},
                    {"en": "META_VALUE_EN", "zh": "META_VALUE_ZH"},
                )
            ],
            items=[figure, table],
            base_dir=self.root,
            lang="bilingual",
        )

        self.assertEqual(returned, output)
        self.assertTrue(output.is_file(), "the output parent directory should be created")
        text = _document_text(output)
        expected_tokens = {
            "REPORT_TITLE_EN", "REPORT_TITLE_ZH", "SUBTITLE_EN", "SUBTITLE_ZH",
            "META_LABEL_EN", "META_LABEL_ZH", "META_VALUE_EN", "META_VALUE_ZH",
            "FIGURE_TITLE_EN", "FIGURE_TITLE_ZH", "CLAIM_EN", "CLAIM_ZH",
            "METHOD_EN", "METHOD_ZH", "RESULT_EN", "RESULT_ZH",
            "INTERPRET_EN", "INTERPRET_ZH", "CAPTION_EN", "CAPTION_ZH",
            "ANNOTATION_KEY_EN", "ANNOTATION_VALUE_EN", "ANNOTATION_KEY_ZH",
            "ANNOTATION_VALUE_ZH", "READ_EN", "READ_ZH", "MISREAD_EN",
            "MISREAD_ZH", "LIMIT_EN", "LIMIT_ZH", "CITATION_EN", "CITATION_ZH",
            "REPRO_EN", "REPRO_ZH", "TABLE_HEADING_EN", "TABLE_HEADING_ZH",
            "TABLE_TITLE_EN", "TABLE_TITLE_ZH", "TABLE_NOTE_EN", "TABLE_NOTE_ZH",
        }
        for token in expected_tokens:
            with self.subTest(token=token):
                self.assertIn(token, text)

    def test_monolingual_mode_keeps_requested_language(self):
        output = self.root / "english.docx"
        build_report(
            output,
            title={"en": "TITLE_EN", "zh": "TITLE_ZH"},
            items=[self._figure(claim={"en": "CLAIM_EN", "zh": "CLAIM_ZH"})],
            base_dir=self.root,
            lang="en",
        )
        text = _document_text(output)
        self.assertIn("TITLE_EN", text)
        self.assertIn("CLAIM_EN", text)
        self.assertNotIn("TITLE_ZH", text)
        self.assertNotIn("CLAIM_ZH", text)

    def test_legacy_caption_fields_and_shared_fields_remain_supported(self):
        output = self.root / "legacy.docx"
        build_report(
            output,
            title="Shared report title",
            items=[
                self._figure(
                    title="Shared figure title",
                    caption=None,
                    caption_en="LEGACY_CAPTION_EN",
                    caption_zh="LEGACY_CAPTION_ZH",
                    annotations=[("Shared key", "Shared value")],
                )
            ],
            base_dir=self.root,
        )
        text = _document_text(output)
        self.assertIn("Shared report title", text)
        self.assertIn("Shared figure title", text)
        self.assertIn("LEGACY_CAPTION_EN", text)
        self.assertIn("LEGACY_CAPTION_ZH", text)
        self.assertIn("Shared key", text)
        self.assertIn("Shared value", text)

    def test_missing_figure_and_csv_files_raise(self):
        cases = [
            self._figure(image="missing.png"),
            self._table(csv="missing.csv"),
        ]
        for item in cases:
            with self.subTest(kind=item["kind"]):
                with self.assertRaises(FileNotFoundError):
                    build_report(
                        self.root / f"{item['kind']}.docx",
                        title="Report",
                        items=[item],
                        base_dir=self.root,
                    )

    def test_missing_required_asset_keys_raise(self):
        cases = [
            {"kind": "figure", "number": "1"},
            {"kind": "table", "number": "1"},
        ]
        for item in cases:
            with self.subTest(kind=item["kind"]):
                with self.assertRaisesRegex(ValueError, "is required"):
                    build_report(self.root / "report.docx", title="Report", items=[item])

    def test_basic_item_schema_is_validated(self):
        cases = [
            ({"kind": "chart", "number": "1", "image": self.image.name}, "kind"),
            ({"kind": [], "number": "1", "image": self.image.name}, "kind"),
            ({"kind": "figure", "image": self.image.name}, "number"),
            (self._figure(status="ready"), "status"),
            (self._figure(status={"en": "final"}), "status"),
            (self._figure(width_in=0), "width_in"),
        ]
        for item, field in cases:
            with self.subTest(field=field):
                with self.assertRaisesRegex(ValueError, field):
                    build_report(
                        self.root / "invalid.docx",
                        title="Report",
                        items=[item],
                        base_dir=self.root,
                    )

    def test_needs_author_decision_is_a_valid_review_status(self):
        output = self.root / "needs-decision.docx"
        build_report(
            output,
            title="Report",
            items=[self._figure(status="needs-author-decision")],
            base_dir=self.root,
        )
        self.assertIn("needs-author-decision", _document_text(output))

    def test_unreadable_assets_raise_contextual_errors(self):
        broken_image = self.root / "broken.png"
        broken_image.write_text("not an image", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "Could not embed image"):
            build_report(
                self.root / "broken-image.docx",
                title="Report",
                items=[self._figure(image=broken_image.name)],
                base_dir=self.root,
            )

        empty_csv = self.root / "empty.csv"
        empty_csv.write_text("", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "Could not read CSV"):
            build_report(
                self.root / "broken-table.docx",
                title="Report",
                items=[self._table(csv=empty_csv.name)],
                base_dir=self.root,
            )


if __name__ == "__main__":
    unittest.main()
